from docx import Document

doc = Document()

doc.add_heading('Повестка в военкомат', 0)

table = doc.add_table(rows=2, cols=2)
table.rows[0].cells[0].text = 'Гражданину '
table.rows[0].cells[1].text = '{{ initials }}'
table.rows[1].cells[0].text = 'Проживающему '
table.rows[1].cells[1].text = '{{ address }}'

p = doc.add_paragraph('')
p.add_run(
    'На основании Федерального Закона '),
p.add_run('"О воинской обязанности и военной службе "').italic=True
p.add_run('Вы подлежите призыву на военную служду и обязаны {{ date }} к {{ time }} час. явиться по адресу: {{ military_address }} для ')
p.add_run('проведения мероприятий связанных с призывом на военную службу.').bold = True

doc.add_heading('При себе иметь: ', 1)
doc.add_paragraph('Паспорт (документ, удостоверяющий личность)',
                       style='List Bullet')
doc.add_paragraph('Ручку',
                       style='List Bullet')
doc.add_paragraph('Медицинскую маску',
                       style='List Bullet')

p = doc.add_paragraph('Подпись {{ signature }}')

p = doc.add_paragraph('Печать {{ stamp }}')

doc.save('lab2_template.docx')
